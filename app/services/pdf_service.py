"""PDF to Word conversion using PyMuPDF + python-docx.

Text PDFs: PyMuPDF block-level text extraction + table detection + header/footer filtering.
Heading detection: text with font size >= 1.3x body size becomes a heading.

For scanned/image-based PDFs: OCR → LLM restructure → markdown → .docx.
"""

import io
import re
from collections import Counter

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement


async def get_pdf_page_count(pdf_file_bytes: bytes) -> int:
    """Get number of pages in PDF."""
    pdf_stream = io.BytesIO(pdf_file_bytes)
    doc = fitz.open(stream=pdf_stream, filetype="pdf")
    count = len(doc)
    doc.close()
    return count


def is_scanned_pdf(file_bytes: bytes) -> bool:
    """Check if any page in the PDF lacks a text layer (scanned/image-based).

    Returns True if ANY page has insufficient extractable text (< 20 chars),
    indicating a scanned/image-based PDF that needs OCR.
    """
    pdf_stream = io.BytesIO(file_bytes)
    doc = fitz.open(stream=pdf_stream, filetype="pdf")
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text().strip()
        # If any page has very little text, treat as scanned
        if len(text) < 20:
            doc.close()
            return True
    doc.close()
    return False


def _estimate_body_font_size(all_blocks: list) -> float:
    """Find the most common font size (mode) across all blocks."""
    sizes = []
    for blk in all_blocks:
        for li in blk.get("lines", []):
            sizes.append(round(li["font_size"]))
    if not sizes:
        return 11.0
    counter = Counter(sizes)
    return float(counter.most_common(1)[0][0])


def _extract_blocks_from_page(page) -> list[dict]:
    """Extract text blocks from a page using PyMuPDF's built-in block grouping.

    PyMuPDF already groups related text into logical blocks — far more reliable
    than line-by-line heuristic merging. Each block typically corresponds to
    a paragraph, heading, or other logical unit.
    """
    blocks = []
    try:
        page_dict = page.get_text("dict")
        for raw_block in page_dict["blocks"]:
            if raw_block["type"] != 0:  # Skip image blocks
                continue

            block_bbox = raw_block.get("bbox", (0, 0, 0, 0))
            lines = []
            all_font_sizes = []

            for line in raw_block.get("lines", []):
                if not line.get("spans"):
                    continue
                first_span = line["spans"][0]
                text = "".join(s["text"] for s in line["spans"])
                if not text.strip():
                    continue

                font_size = first_span.get("size", 11)
                font_name = first_span.get("font", "")
                is_bold = "Bold" in font_name
                is_italic = "Italic" in font_name or "Oblique" in font_name
                line_bbox = line.get("bbox", block_bbox)

                lines.append({
                    "text": text.strip(),
                    "font_size": font_size,
                    "font_name": font_name,
                    "is_bold": is_bold,
                    "is_italic": is_italic,
                    "x0": line_bbox[0],
                    "y0": line_bbox[1],
                    "x1": line_bbox[2],
                    "y1": line_bbox[3],
                })
                all_font_sizes.append(font_size)

            if not lines:
                continue

            avg_font_size = sum(all_font_sizes) / len(all_font_sizes)
            max_font_size = max(all_font_sizes)

            blocks.append({
                "lines": lines,
                "bbox": block_bbox,
                "avg_font_size": avg_font_size,
                "max_font_size": max_font_size,
            })
    except Exception:
        pass

    return blocks


def _extract_tables_from_page(page) -> list[dict]:
    """Detect tables on a page using PyMuPDF's built-in table detection."""
    tables = []
    try:
        found = page.find_tables()
        if found and found.tables:
            for table in found.tables:
                cells = table.extract()
                if not cells or all(not any(c for c in row) for row in cells):
                    continue
                tables.append({
                    "bbox": table.bbox,
                    "rows": cells,
                })
    except Exception:
        pass
    return tables


def _bbox_overlaps(bbox1, bbox2, threshold: float = 0.5) -> bool:
    """Check if bbox1 significantly overlaps with bbox2."""
    x0_1, y0_1, x1_1, y1_1 = bbox1
    x0_2, y0_2, x1_2, y1_2 = bbox2

    ix0 = max(x0_1, x0_2)
    iy0 = max(y0_1, y0_2)
    ix1 = min(x1_1, x1_2)
    iy1 = min(y1_1, y1_2)

    if ix0 >= ix1 or iy0 >= iy1:
        return False

    intersection = (ix1 - ix0) * (iy1 - iy0)
    area1 = (x1_1 - x0_1) * (y1_1 - y0_1)
    return intersection / max(area1, 1) > threshold


def _is_header_footer(bbox, page_height: float, margin_ratio: float = 0.15) -> bool:
    """Check if a text block is in the header or footer zone."""
    y0, y1 = bbox[1], bbox[3]
    return y1 < page_height * margin_ratio or y0 > page_height * (1 - margin_ratio)


_LIST_PATTERNS = [
    re.compile(r'^\s*[\-\•\*\▪\▸\➤\–\—]\s'),
    re.compile(r'^\s*\d+[\.\)]\s'),
    re.compile(r'^\s*[a-zA-Z][\.\)]\s'),
    re.compile(r'^\s*\([a-zA-Z\d]+\)\s'),
]


def _is_list_item(text: str) -> bool:
    """Check if text looks like a bullet or numbered list item."""
    return any(p.match(text) for p in _LIST_PATTERNS)


def _apply_run_styling(run, line_info: dict):
    """Apply font styling to a run."""
    run.font.size = Pt(line_info["font_size"])
    run.bold = line_info["is_bold"]
    run.italic = line_info["is_italic"]
    if line_info["font_name"]:
        run.font.name = line_info["font_name"]
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = rPr.makeelement(qn("w:rFonts"), {})
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), line_info["font_name"])
        rFonts.set(qn("w:hAnsi"), line_info["font_name"])


def _set_paragraph_spacing(p, space_before=0, space_after=6):
    """Set paragraph spacing in twips."""
    pPr = p._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = pPr.makeelement(qn("w:spacing"), {})
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(int(space_before * 20)))
    spacing.set(qn("w:after"), str(int(space_after * 20)))


def _add_formatted_paragraph(docx, lines: list, body_font_size: float):
    """Add a paragraph to the docx with proper formatting.

    Detects headings, list items, and regular paragraphs.
    """
    joined_text = " ".join(l["text"] for l in lines)
    is_heading = any(l["font_size"] >= body_font_size * 1.3 for l in lines)
    is_list = len(lines) == 1 and _is_list_item(lines[0]["text"])

    p = docx.add_paragraph()

    if is_heading:
        heading_size = max(l["font_size"] for l in lines)
        heading_size = min(heading_size, 28)
        run = p.add_run(joined_text)
        run.font.size = Pt(heading_size)
        run.bold = True
        _set_paragraph_spacing(p, space_before=12, space_after=6)
    elif is_list:
        # Use single-line text as-is, preserving the bullet/number prefix
        for i, line_info in enumerate(lines):
            separator = "" if i == 0 else " "
            run = p.add_run(separator + line_info["text"])
            _apply_run_styling(run, line_info)
        _set_paragraph_spacing(p, space_before=0, space_after=2)
        p.paragraph_format.left_indent = Inches(0.25)
    else:
        for i, line_info in enumerate(lines):
            separator = "" if i == 0 else " "
            run = p.add_run(separator + line_info["text"])
            _apply_run_styling(run, line_info)
        _set_paragraph_spacing(p, space_before=0, space_after=6)


def _add_table_to_docx(docx, rows: list[list[str]]):
    """Build a docx table from PyMuPDF table extraction output."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    num_rows = len(rows)
    table = docx.add_table(rows=num_rows, cols=num_cols, style="Table Grid")
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(cell_text or ""))
                run.font.size = Pt(9)
                if r_idx == 0:
                    run.bold = True
                _set_paragraph_spacing(p, space_before=0, space_after=0)
    docx.add_paragraph()  # spacing after table


async def convert_pdf_to_word(pdf_file_bytes: bytes, filename: str) -> bytes:
    """Convert PDF to DOCX using PyMuPDF block-level text extraction + table detection.

    Uses PyMuPDF's built-in block grouping for reliable paragraph detection,
    find_tables() for table extraction, and filters out headers/footers.

    Raises ValueError if no text can be extracted (e.g. scanned/image PDF).
    """
    pdf_stream = io.BytesIO(pdf_file_bytes)
    pdf = fitz.open(stream=pdf_stream, filetype="pdf")

    # Pass 1: collect all blocks and tables
    all_pages_blocks = []
    all_pages_tables = []
    page_heights = []

    for page_num in range(len(pdf)):
        page = pdf[page_num]
        page_rect = page.rect
        page_heights.append(page_rect.height)

        blocks = _extract_blocks_from_page(page)
        tables = _extract_tables_from_page(page)
        all_pages_blocks.append(blocks)
        all_pages_tables.append(tables)

    # Estimate body font size from all blocks
    all_blocks_flat = [blk for page_blocks in all_pages_blocks for blk in page_blocks]
    body_font_size = _estimate_body_font_size(all_blocks_flat)

    docx = Document()

    # Default style
    style = docx.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(body_font_size)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    total_elements = 0

    for page_num in range(len(pdf)):
        blocks = all_pages_blocks[page_num]
        tables = all_pages_tables[page_num]
        page_h = page_heights[page_num]

        # Mark blocks that overlap with tables (to avoid duplicate text)
        overlapped_indices = set()
        if tables:
            for ti, table in enumerate(tables):
                for bi, blk in enumerate(blocks):
                    if _bbox_overlaps(blk["bbox"], table["bbox"]):
                        overlapped_indices.add(bi)

        # Sort blocks by Y position (top to bottom)
        blocks_with_idx = [(bi, blk) for bi, blk in enumerate(blocks)]
        blocks_with_idx.sort(key=lambda x: x[1]["bbox"][1])

        # Interleave tables and text blocks by Y position
        elements = []

        # Build list of (y_position, type, data)
        for bi, blk in blocks_with_idx:
            if bi in overlapped_indices:
                continue
            # Skip header/footer blocks (but keep if they're the only content)
            if len(blocks) > 3 and _is_header_footer(blk["bbox"], page_h):
                continue
            elements.append((blk["bbox"][1], "text", blk))

        for table in tables:
            elements.append((table["bbox"][1], "table", table))

        # Sort all elements by Y position
        elements.sort(key=lambda x: x[0])

        for _, elem_type, data in elements:
            if elem_type == "table":
                _add_table_to_docx(docx, data["rows"])
                total_elements += 1
            elif elem_type == "text":
                _add_formatted_paragraph(docx, data["lines"], body_font_size)
                total_elements += 1

        # Page break between pages (except last)
        if page_num < len(all_pages_blocks) - 1:
            docx.add_page_break()

    pdf.close()

    if total_elements == 0:
        raise ValueError(
            "This PDF appears to be a scanned document or image without a text layer. "
            "PDF to Word conversion requires PDFs with selectable text. "
            "Please use an OCR tool first to add a text layer to your PDF."
        )

    buf = io.BytesIO()
    docx.save(buf)
    buf.seek(0)
    return buf.read()


def markdown_to_docx(markdown_text: str) -> bytes:
    """Convert structured markdown (from LLM) to .docx bytes.

    Handles: headings (# ## ###), tables (|...|), bullet lists (- *),
    numbered lists (1. 2.), block quotes (>), and regular paragraphs.
    Inline formatting: **bold** and *italic* supported.
    """
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    lines = markdown_text.split("\n")
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Empty line: flush pending table, then skip
        if not line:
            if in_table and table_rows:
                _build_table(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        # Table row detection
        if line.startswith("|") and line.endswith("|"):
            # Skip separator rows like |---|---|
            if re.match(r"^\|[\s\-:|]+\|$", line):
                in_table = True
                i += 1
                continue
            cells = [c.strip() for c in line.split("|")[1:-1]]
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        else:
            # Flush pending table
            if in_table and table_rows:
                _build_table(doc, table_rows)
                table_rows = []
                in_table = False

        # Heading: ###, ##, #
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            p = doc.add_paragraph()
            p.style = doc.styles[f"Heading {level}"]
            _add_inline_runs(p, text)
            i += 1
            continue

        # Bullet list: "- " or "* "
        if re.match(r"^[\-\*]\s+", line):
            text = re.sub(r"^[\-\*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, text)
            i += 1
            continue

        # Numbered list: "1. " etc.
        num_match = re.match(r"^(\d+)\.\s+(.+)$", line)
        if num_match:
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, num_match.group(2))
            i += 1
            continue

        # Block quote: "> "
        if line.startswith("> "):
            text = line[2:]
            p = doc.add_paragraph()
            _add_inline_runs(p, text)
            p.paragraph_format.left_indent = Inches(0.5)
            _set_paragraph_spacing(p, space_before=0, space_after=6)
            i += 1
            continue

        # Regular paragraph
        _add_inline_runs(doc.add_paragraph(), line)
        i += 1

    # Flush remaining table
    if in_table and table_rows:
        _build_table(doc, table_rows)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_inline_runs(paragraph, text: str):
    """Add text with **bold** and *italic* markdown parsing to a paragraph."""
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*)")
    last_end = 0
    for match in pattern.finditer(text):
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            run.italic = True
        last_end = match.end()
    if last_end < len(text):
        paragraph.add_run(text[last_end:])
    _set_paragraph_spacing(paragraph, space_before=0, space_after=6)


def _build_table(doc, rows):
    """Build a docx table from a list of cell lists."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols, style="Table Grid")
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                _add_inline_runs(p, cell_text)
                # Bold first row as header
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True
    doc.add_paragraph()  # spacing after table


async def convert_scanned_pdf_to_word(
    pdf_file_bytes: bytes,
    filename: str,
    ocr_text: str,
    restructured_markdown: str,
    page_images: list[bytes] | None = None,
) -> bytes:
    """Convert scanned PDF to Word using OCR + LLM markdown, with image fallback.

    Prefers structured markdown (from LLM restructure) over raw OCR text.
    Falls back to image embedding only when no usable text is available.
    """
    source = restructured_markdown.strip() or ocr_text.strip()
    if source:
        return markdown_to_docx(source)
    # No usable text: embed page images as last resort
    return _images_to_docx(page_images or [], filename)


def _images_to_docx(page_images: list[bytes], filename: str = "document") -> bytes:
    """Create a .docx with each page image embedded as a full-page picture."""
    import io as _io
    from docx.shared import Inches as _Inches
    from docx.enum.section import WD_ORIENT

    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = _Inches(0.5)
        section.bottom_margin = _Inches(0.5)
        section.left_margin = _Inches(0.5)
        section.right_margin = _Inches(0.5)

    for i, img_bytes in enumerate(page_images):
        if i > 0:
            doc.add_page_break()

        # Add page label
        p = doc.add_paragraph()
        run = p.add_run(f"Page {i + 1}")
        run.bold = True
        run.font.size = Pt(11)
        _set_paragraph_spacing(p, space_before=0, space_after=6)

        # Embed image
        img_stream = _io.BytesIO(img_bytes)
        try:
            doc.add_picture(img_stream, width=_Inches(6.0))
        except Exception:
            # If image embedding fails, add placeholder
            p2 = doc.add_paragraph()
            p2.add_run(f"[Image for page {i + 1} could not be embedded]")

        # Add separator
        if i < len(page_images) - 1:
            p_sep = doc.add_paragraph()
            p_sep.add_run("—" * 40)
            p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
